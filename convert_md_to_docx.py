#!/usr/bin/env python3
import argparse
from docx import Document
from docx.shared import Pt

from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_toc(paragraph):
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)

    run = paragraph.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._r.append(instrText)

    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldChar)

    run = paragraph.add_run('Table of Contents (Right-click and "Update Field" in Word)')
    
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)

def is_heading(line):
    s = line.lstrip()
    if not s.startswith('#'):
        return 0, None
    hashes = 0
    for ch in s:
        if ch == '#':
            hashes += 1
        else:
            break
    text = s[hashes:].strip()
    return hashes, text

def convert_md_to_docx(md_text, out_path):
    doc = Document()
    
    # Add TOC Placeholder at the top
    doc.add_heading('Table of Contents', level=1)
    toc_para = doc.add_paragraph()
    add_toc(toc_para)
    doc.add_page_break()

    in_code = False
    code_lines = []
    
    # Flag to skip the manual TOC in the MD file if it exists
    skip_manual_toc = False
    
    for raw_line in md_text.splitlines():
        line = raw_line.rstrip('\n')
        
        # Skip everything until first real section after '## Contents'
        if line.strip() == '## Contents' or line.strip() == '# Contents':
            skip_manual_toc = True
            continue
        
        if skip_manual_toc:
            if line.startswith('#'):
                skip_manual_toc = False # Stop skipping when we hit next header
            else:
                continue # Skip the list items of manual TOC

        if line.strip().startswith('```'):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                in_code = False
                p = doc.add_paragraph()
                run = p.add_run('\n'.join(code_lines))
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
            continue
        
        if in_code:
            code_lines.append(line)
            continue

        if not line.strip():
            doc.add_paragraph('')
            continue

        level, text = is_heading(line)
        if level:
            doc.add_heading(text, level=min(level, 9))
            continue

        stripped = line.lstrip()
        if stripped.startswith(('- ', '* ')):
            content = stripped[2:].strip()
            doc.add_paragraph(content, style='List Bullet')
            continue

        if set(line.strip()) == set('-') and len(line.strip()) >= 3:
            doc.add_paragraph('')
            continue

        doc.add_paragraph(line)

    doc.save(out_path)


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Convert Markdown to simple DOCX')
    parser.add_argument('input_md', help='Path to input Markdown file')
    parser.add_argument('output_docx', help='Path to output DOCX file')
    args = parser.parse_args()

    with open(args.input_md, 'r', encoding='utf-8') as f:
        md = f.read()

    convert_md_to_docx(md, args.output_docx)
